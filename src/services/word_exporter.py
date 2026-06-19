import io
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from sqlalchemy.orm import Session, joinedload
from src.models.entities import PlanMacro, Policy, StrategicItem, Activity, Task, Evidence, Responsible, User

class WordExporterService:
    """
    Servicio encargado de generar informes oficiales en formato Word (.docx)
    con el diseño, estilo y estructura del INFORME DE GESTIÓN ANUAL 2025.
    """
    
    @staticmethod
    def _set_cell_background(cell, fill_hex):
        """Aplica color de fondo a una celda."""
        tcPr = cell._tc.get_or_add_tcPr()
        shd_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
        shd = parse_xml(shd_xml)
        tcPr.append(shd)

    @staticmethod
    def _set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        """Establece márgenes internos (padding) para una celda en dxa."""
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    @staticmethod
    def _add_styled_heading(doc, text, level, space_before=12, space_after=6):
        """Agrega un encabezado con tipografía Arial y colores institucionales."""
        heading = doc.add_heading(text, level=level)
        heading.paragraph_format.space_before = Pt(space_before)
        heading.paragraph_format.space_after = Pt(space_after)
        heading.paragraph_format.keep_with_next = True
        
        # Color primario institucional verde (#00594e) para nivel 1 y 2, y azul (#0f766e) para nivel 3 y 4
        color_hex = "00594e" if level in [1, 2] else "0f766e"
        r = int(color_hex[0:2], 16)
        g = int(color_hex[2:4], 16)
        b = int(color_hex[4:6], 16)
        
        for run in heading.runs:
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(r, g, b)
            run.font.bold = True
            
        return heading

    @staticmethod
    def generate_report(db: Session, start_date: datetime, end_date: datetime, filter_label: str) -> io.BytesIO:
        """
        Genera el documento Word siguiendo la estructura del Informe de Gestión Anual 2025.
        """
        doc = Document()
        
        # Determinar el tipo de informe según el filter_label
        label_lower = filter_label.lower()
        if "año" in label_lower or "vigencia" in label_lower or "completo" in label_lower:
            report_type_title = "INFORME DE GESTIÓN ANUAL"
        elif "semestre" in label_lower:
            report_type_title = "INFORME DE GESTIÓN SEMESTRAL"
        elif "trimestre" in label_lower or "q1" in label_lower or "q2" in label_lower or "q3" in label_lower or "q4" in label_lower:
            report_type_title = "INFORME DE GESTIÓN TRIMESTRAL"
        elif "mes" in label_lower:
            report_type_title = "INFORME DE GESTIÓN MENSUAL"
        else:
            report_type_title = "INFORME DE GESTIÓN"

        # Logo path
        logo_path = "assets/images/logo.png"

        # --- CONFIGURACIÓN DE PÁGINA (Márgenes de 1 pulgada) ---
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
            # Cabecera/pie diferente para primera página (portada)
            section.different_first_page_header_footer = True
            
            # Configurar cabecera de páginas internas con tabla para logo + texto
            header = section.header
            for p in header.paragraphs:
                p.text = ""
            
            # Tabla de cabecera
            header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
            header_table.style = 'Normal Table'
            hdr_cells = header_table.rows[0].cells
            
            # Logo a la izquierda
            p_logo_hdr = hdr_cells[0].paragraphs[0]
            p_logo_hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if os.path.exists(logo_path):
                p_logo_hdr.add_run().add_picture(logo_path, width=Inches(1.1))
            else:
                p_logo_hdr.add_run("UNITRÓPICO").bold = True
                p_logo_hdr.runs[0].font.size = Pt(9.5)
                
            # Metadatos a la derecha
            p_meta_hdr = hdr_cells[1].paragraphs[0]
            p_meta_hdr.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_meta_hdr.text = f"Gestión de Talento Humano\n{report_type_title}\nPeriodo: {filter_label}"
            for run in p_meta_hdr.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8.0)
                run.font.color.rgb = RGBColor(100, 116, 139)

            # Agregar una línea separadora verde debajo de la tabla de cabecera
            p_sep = header.add_paragraph()
            p_sep.paragraph_format.space_before = Pt(2)
            p_sep.paragraph_format.space_after = Pt(0)
            run_sep = p_sep.add_run("―" * 55)
            run_sep.font.color.rgb = RGBColor(0, 89, 78) # Verde
            run_sep.font.size = Pt(9)
            
            # Configurar pie de páginas internas
            footer = section.footer
            for p in footer.paragraphs:
                p.text = ""
                
            footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
            footer_table.style = 'Normal Table'
            ft_cells = footer_table.rows[0].cells
            
            # Texto a la izquierda
            p_left = ft_cells[0].paragraphs[0]
            p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p_left.text = "Universidad Internacional del Trópico Americano"
            for run in p_left.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8.0)
                run.font.color.rgb = RGBColor(100, 116, 139)
                
            # Texto a la derecha
            p_right = ft_cells[1].paragraphs[0]
            p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_right.text = "Oficina de Talento Humano  |  Confidencial"
            for run in p_right.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(8.0)
                run.font.color.rgb = RGBColor(100, 116, 139)

        # Configurar estilos por defecto (Arial, 11pt, color corporativo gris oscuro)
        style_normal = doc.styles['Normal']
        style_normal.font.name = 'Arial'
        style_normal.font.size = Pt(11)
        style_normal.font.color.rgb = RGBColor(51, 65, 85)
        style_normal.paragraph_format.line_spacing = 1.15
        style_normal.paragraph_format.space_after = Pt(6)

        # --- DATA GATHERING ---
        target_year = start_date.year
        macro = db.query(PlanMacro).filter(PlanMacro.year == target_year).options(
            joinedload(PlanMacro.policies)
            .joinedload(Policy.strategic_items)
            .joinedload(StrategicItem.activities)
            .joinedload(Activity.tasks)
            .joinedload(Task.responsibles)
        ).first()

        # Fallback si no hay año coincidente
        if not macro:
            macro = db.query(PlanMacro).order_by(PlanMacro.year.desc()).first()

        # --- PORTADA (ESTILO INFORME 2025) ---
        p_logo = doc.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if os.path.exists(logo_path):
            p_logo.add_run().add_picture(logo_path, width=Inches(1.8))
        else:
            p_logo.add_run("[UNITRÓPICO]").bold = True
        
        doc.add_paragraph("\n" * 4)
        
        # Título principal de la portada
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_title.add_run(report_type_title)
        run_title.font.size = Pt(28)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(0, 89, 78) # Verde institucional
        
        # Subtítulo (Periodo y Vigencia)
        p_year = doc.add_paragraph()
        p_year.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_year = p_year.add_run(f"Periodo: {filter_label}")
        run_year.font.size = Pt(18)
        run_year.font.bold = True
        run_year.font.color.rgb = RGBColor(100, 116, 139)
        
        doc.add_paragraph("\n" * 5)
        
        # Bloque de metadatos de autoría
        meta_data = [
            ("Elaborado por:", "Lilia Andrea Nocua Neme"),
            ("Cargo:", "Jefe Oficina de Talento Humano"),
            ("Fecha de generación:", datetime.now().strftime("%d de %B de %Y")),
            ("Procesos Relacionados:", "Todos los procesos operativos y estratégicos de TH"),
            ("Versión del reporte:", "01")
        ]
        
        for label, val in meta_data:
            p_meta = doc.add_paragraph()
            run_lbl = p_meta.add_run(f"{label} ")
            run_lbl.bold = True
            run_lbl.font.size = Pt(10)
            run_val = p_meta.add_run(val)
            run_val.font.size = Pt(10)
            p_meta.paragraph_format.space_after = Pt(2)
            
        doc.add_page_break()

        # --- INTRODUCCIÓN ---
        WordExporterService._add_styled_heading(doc, "INTRODUCCIÓN", level=1, space_before=0, space_after=12)
        doc.add_paragraph(
            "La Oficina de Talento Humano, en cumplimiento de su misión de dirigir, coordinar y controlar la gestión "
            "estratégica del talento humano al servicio de la Universidad Internacional del Trópico Americano - Unitrópico, "
            "presenta este Informe de Gestión Anual. Durante este periodo, se consolidaron acciones orientadas a garantizar la "
            "selección, vinculación, formación, evaluación, bienestar y retiro de los servidores administrativos y profesores, "
            "fortaleciendo su compromiso con la misión, los objetivos y metas institucionales en la construcción de entornos "
            "laborales colaborativos, seguros y productivos."
        )
        doc.add_paragraph(
            "El presente informe integra los resultados de la gestión desarrollada, estructurándose en las secciones de "
            "cifras globales e hitos alcanzados, avance en las Políticas Institucionales de Gestión Estratégica de Talento "
            "Humano, Estímulos e Integridad, avance de las metas del Plan de Desarrollo Institucional, y firmas de aprobación. "
            "Este balance de gestión evidencia los logros cuantitativos obtenidos y el cumplimiento de las metas del plan operativo."
        )
        
        doc.add_page_break()

        # --- SECCIÓN 1: CIFRAS E HITOS ---
        WordExporterService._add_styled_heading(doc, "1. CIFRAS E HITOS", level=1, space_before=12, space_after=12)
        
        # 1.1 Cifras en la gestión
        WordExporterService._add_styled_heading(doc, "1.1 Cifras en la gestión de la Oficina de Talento Humano", level=2, space_before=6, space_after=8)
        doc.add_paragraph(
            "A continuación, se presenta una visión consolidada de los principales indicadores cuantitativos que reflejan el alcance, "
            "la dinámica y el impacto de la gestión desarrollada por la Oficina de Talento Humano durante la vigencia reportada:"
        )

        # Consultas de cifras del sistema
        total_tasks = db.query(Task).count()
        completed_tasks = db.query(Task).filter(Task.status == "Cumplida").count()
        in_process_tasks = db.query(Task).filter(Task.status == "En Proceso").count()
        pending_tasks = db.query(Task).filter(Task.status == "Pendiente").count()
        total_evidences = db.query(Evidence).count()
        total_responsibles = db.query(Responsible).count()
        avg_progress = macro.progress if macro else 0.0

        # Tabla de cifras
        cifras_data = [
            ("Total de Tareas Programadas", str(total_tasks)),
            ("Tareas Cumplidas Exitosamente", str(completed_tasks)),
            ("Tareas en Ejecución / Proceso", str(in_process_tasks)),
            ("Tareas Pendientes", str(pending_tasks)),
            ("Porcentaje de Cumplimiento Global", f"{avg_progress:.1f}%"),
            ("Evidencias y Soportes Cargados", str(total_evidences)),
            ("Funcionarios Responsables Vinculados", str(total_responsibles))
        ]

        table_cifras = doc.add_table(rows=1, cols=2)
        table_cifras.style = 'Table Grid'
        
        # Cabecera
        hdr_cells = table_cifras.rows[0].cells
        hdr_cells[0].text = "INDICADOR / ACCIÓN DE GESTIÓN"
        hdr_cells[1].text = "TOTAL REGISTRADO"
        WordExporterService._set_cell_background(hdr_cells[0], "00594e")
        WordExporterService._set_cell_background(hdr_cells[1], "00594e")
        for cell in hdr_cells:
            WordExporterService._set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)

        # Rellenar datos
        for idx, (ind, val) in enumerate(cifras_data):
            row_cells = table_cifras.add_row().cells
            row_cells[0].text = ind
            row_cells[1].text = val
            
            # Zebra rows background
            bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
            WordExporterService._set_cell_background(row_cells[0], bg_color)
            WordExporterService._set_cell_background(row_cells[1], bg_color)
            
            WordExporterService._set_cell_margins(row_cells[0], top=80, bottom=80, left=150, right=150)
            WordExporterService._set_cell_margins(row_cells[1], top=80, bottom=80, left=150, right=150)
            
            # Align right for numbers
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            row_cells[1].paragraphs[0].runs[0].font.bold = True

        # Ajustar anchos de columnas
        col_widths = [Inches(4.5), Inches(2.0)]
        for row in table_cifras.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width

        doc.add_paragraph("\n")

        # 1.2 Hitos en la gestión
        WordExporterService._add_styled_heading(doc, "1.2 Hitos en la gestión de la Oficina de Talento Humano", level=2, space_before=6, space_after=8)
        doc.add_paragraph(
            "A continuación se destacan los logros estratégicos y metas alcanzadas de alto impacto registrados en el sistema:"
        )

        hitos_list = [
            f"Consolidación y cumplimiento del {avg_progress:.1f}% de avance promedio en la ejecución de las metas estratégicas de Talento Humano.",
            f"Se registraron y documentaron un total de {completed_tasks} tareas completadas con sus respectivos soportes de evidencia digitalizados.",
            f"Seguimiento y control de {total_evidences} evidencias que avalan la transparencia y veracidad de los soportes cargados en la plataforma.",
            "Normalización y distribución equilibrada de pesos ponderados en toda la estructura del Plan de Acción TH.",
            "Puesta en marcha del sistema integral con control de accesos seguro por roles y trazabilidad de cambios por usuario (ChangeLog)."
        ]

        for h in hitos_list:
            doc.add_paragraph(h, style='List Bullet')

        doc.add_page_break()

        # --- SECCIÓN 2: AVANCE EN POLÍTICAS INSTITUCIONALES ---
        WordExporterService._add_styled_heading(doc, "2. AVANCE EN LAS POLÍTICAS INSTITUCIONALES", level=1, space_before=12, space_after=12)
        doc.add_paragraph(
            "Las actividades desarrolladas se estructuran de manera sistemática bajo las Políticas Institucionales de Gestión "
            "Estratégica de Talento Humano, Estímulos e Incentivos, e Integridad. A continuación, se detalla el cumplimiento de cada una:"
        )

        if not macro or not macro.policies:
            doc.add_paragraph("No hay políticas registradas en este plan macro.")
        else:
            # Filtrar solo Políticas 1, 2 y 3 (excluyendo Plan de Desarrollo que va en Sección 3)
            sec2_policies = [p for p in macro.policies if p.name != "Plan de Desarrollo Institucional"]
            
            for pol_idx, pol in enumerate(sec2_policies, 1):
                WordExporterService._add_styled_heading(doc, f"2.{pol_idx} Política de {pol.name}", level=2, space_before=10, space_after=8)
                
                # Resumen de la política
                doc.add_paragraph(
                    f"La política de {pol.name} representa el {pol.weight:.1f}% del peso ponderado global del Plan Maestro. "
                    f"Para la vigencia reportada, registra un avance consolidado del {pol.progress:.1f}%."
                )
                
                # Loop por StrategicItems (Strategies)
                for item_idx, item in enumerate(pol.strategic_items, 1):
                    WordExporterService._add_styled_heading(doc, f"2.{pol_idx}.{item_idx}. {item.name}", level=3, space_before=8, space_after=6)
                    
                    # Filtrar tareas por periodo
                    filtered_tasks = []
                    for act in item.activities:
                        for t in act.tasks:
                            t_start = t.start_date
                            t_end = t.end_date
                            if t_start and t_end:
                                t_start_dt = datetime.combine(t_start, datetime.min.time())
                                t_end_dt = datetime.combine(t_end, datetime.max.time())
                                if t_start_dt <= end_date and t_end_dt >= start_date:
                                    filtered_tasks.append(t)
                            elif t.fulfillment_date:
                                if start_date <= t.fulfillment_date <= end_date:
                                    filtered_tasks.append(t)

                    if not filtered_tasks:
                        doc.add_paragraph("No se registran tareas ejecutadas o programadas para este compromiso en el periodo seleccionado.")
                        continue
                        
                    doc.add_paragraph(
                        f"A continuación se relacionan los avances, comentarios, justificaciones de retraso y enlaces de "
                        f"evidencia correspondientes a las tareas programadas:"
                    )

                    # Crear tabla de tareas
                    table = doc.add_table(rows=1, cols=5)
                    table.style = 'Table Grid'
                    
                    # Configurar cabecera
                    hdr = table.rows[0].cells
                    hdr[0].text = "Tarea / Descripción"
                    hdr[1].text = "Responsable(s)"
                    hdr[2].text = "Avance (%)"
                    hdr[3].text = "Comentarios / Observaciones"
                    hdr[4].text = "Enlace de Evidencia"
                    
                    for c in hdr:
                        WordExporterService._set_cell_background(c, "0f766e")
                        WordExporterService._set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                        for p in c.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(255, 255, 255)
                                r.font.size = Pt(9.5)

                    # Rellenar filas de tareas
                    for r_idx, t in enumerate(filtered_tasks):
                        row = table.add_row().cells
                        
                        # Celda 1: Nombre de la tarea
                        t_lbl = f"{t.name}"
                        if t.start_date and t.end_date:
                            t_lbl += f"\n(Plazo: {t.start_date.strftime('%d/%m')} al {t.end_date.strftime('%d/%m/%y')})"
                        row[0].text = t_lbl
                        
                        # Celda 2: Responsables
                        res_lbl = ", ".join([resp.name for resp in t.responsibles]) if t.responsibles else "Sin asignar"
                        row[1].text = res_lbl
                        
                        # Celda 3: Avance % / Estado
                        row[2].text = f"{t.status}\n({t.progress:.1f}%)"
                        
                        # Celda 4: Comentarios, observaciones y justificación si está vencida
                        comment_text = t.observations or "Sin observaciones registradas."
                        if t.delay_justification:
                            comment_text += f"\n\n⚠️ JUSTIFICACIÓN RETRASO:\n{t.delay_justification}"
                        row[3].text = comment_text
                        
                        # Celda 5: Evidencias con enlaces
                        evidences = db.query(Evidence).filter(Evidence.task_id == t.id).all()
                        p_ev = row[4].paragraphs[0]
                        if evidences:
                            for e_idx, ev in enumerate(evidences, 1):
                                desc = ev.description or f"Evidencia {e_idx}"
                                url = ev.url if ev.url.startswith(("http://", "https://")) else f"https://{ev.url}"
                                run_link = p_ev.add_run(f"[{desc}]\n")
                                run_link.font.underline = True
                                run_link.font.color.rgb = RGBColor(0, 89, 78) # Verde
                                run_link.font.bold = True
                                
                                # Nota del cargador si existe
                                who = f" (por: {ev.uploaded_by_name})" if ev.uploaded_by_name else ""
                                run_sub = p_ev.add_run(f"{url}{who}\n\n")
                                run_sub.font.size = Pt(7.5)
                                run_sub.font.color.rgb = RGBColor(100, 116, 139)
                        else:
                            p_ev.text = "Sin evidencias."
                            
                        # Zebra striping
                        row_bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
                        for cell in row:
                            WordExporterService._set_cell_background(cell, row_bg)
                            WordExporterService._set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                            # Ajustes tipográficos para celdas
                            for p in cell.paragraphs:
                                for r in p.runs:
                                    r.font.size = Pt(8.5)
                                    
                        # Alineaciones específicas
                        row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    # Configurar anchos de columna de tabla de tareas
                    task_col_widths = [Inches(1.8), Inches(1.0), Inches(0.8), Inches(1.8), Inches(1.1)]
                    for r in table.rows:
                        for c_idx, width in enumerate(task_col_widths):
                            r.cells[c_idx].width = width
                            
                    doc.add_paragraph("\n")

        doc.add_page_break()

        # --- SECCIÓN 3: PLAN DE DESARROLLO INSTITUCIONAL ---
        WordExporterService._add_styled_heading(doc, "3. PLAN DE DESARROLLO INSTITUCIONAL", level=1, space_before=12, space_after=12)
        doc.add_paragraph(
            "En el marco del Plan de Desarrollo Institucional, se realiza seguimiento a las metas asociadas al fortalecimiento, "
            "la modernización y el desarrollo estratégico de la planta física y tecnológica de la institución:"
        )

        p_pdi = next((p for p in macro.policies if p.name == "Plan de Desarrollo Institucional"), None)
        if not p_pdi:
            doc.add_paragraph("No se registran metas del Plan de Desarrollo Institucional en este plan macro.")
        else:
            for item_idx, item in enumerate(p_pdi.strategic_items, 1):
                WordExporterService._add_styled_heading(doc, f"3.{item_idx} {item.name}", level=2, space_before=10, space_after=8)
                
                # Filtrar tareas por periodo
                filtered_tasks = []
                for act in item.activities:
                    for t in act.tasks:
                        t_start = t.start_date
                        t_end = t.end_date
                        if t_start and t_end:
                            t_start_dt = datetime.combine(t_start, datetime.min.time())
                            t_end_dt = datetime.combine(t_end, datetime.max.time())
                            if t_start_dt <= end_date and t_end_dt >= start_date:
                                filtered_tasks.append(t)
                        elif t.fulfillment_date:
                            if start_date <= t.fulfillment_date <= end_date:
                                filtered_tasks.append(t)

                if not filtered_tasks:
                    doc.add_paragraph("No se registran tareas de seguimiento para esta meta en el periodo seleccionado.")
                    continue

                # Crear tabla de tareas
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                
                # Configurar cabecera
                hdr = table.rows[0].cells
                hdr[0].text = "Tarea / Descripción"
                hdr[1].text = "Responsable(s)"
                hdr[2].text = "Avance (%)"
                hdr[3].text = "Comentarios / Observaciones"
                hdr[4].text = "Enlace de Evidencia"
                
                for c in hdr:
                    WordExporterService._set_cell_background(c, "0f766e")
                    WordExporterService._set_cell_margins(c, top=80, bottom=80, left=100, right=100)
                    for p in c.paragraphs:
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for r in p.runs:
                            r.font.bold = True
                            r.font.color.rgb = RGBColor(255, 255, 255)
                            r.font.size = Pt(9.5)

                # Rellenar filas de tareas
                for r_idx, t in enumerate(filtered_tasks):
                    row = table.add_row().cells
                    
                    # Celda 1: Nombre de la tarea
                    t_lbl = f"{t.name}"
                    if t.start_date and t.end_date:
                        t_lbl += f"\n(Plazo: {t.start_date.strftime('%d/%m')} al {t.end_date.strftime('%d/%m/%y')})"
                    row[0].text = t_lbl
                    
                    # Celda 2: Responsables
                    res_lbl = ", ".join([resp.name for resp in t.responsibles]) if t.responsibles else "Sin asignar"
                    row[1].text = res_lbl
                    
                    # Celda 3: Avance % / Estado
                    row[2].text = f"{t.status}\n({t.progress:.1f}%)"
                    
                    # Celda 4: Comentarios, observaciones y justificación si está vencida
                    comment_text = t.observations or "Sin observaciones registradas."
                    if t.delay_justification:
                        comment_text += f"\n\n⚠️ JUSTIFICACIÓN RETRASO:\n{t.delay_justification}"
                    row[3].text = comment_text
                    
                    # Celda 5: Evidencias con enlaces
                    evidences = db.query(Evidence).filter(Evidence.task_id == t.id).all()
                    p_ev = row[4].paragraphs[0]
                    if evidences:
                        for e_idx, ev in enumerate(evidences, 1):
                            desc = ev.description or f"Evidencia {e_idx}"
                            url = ev.url if ev.url.startswith(("http://", "https://")) else f"https://{ev.url}"
                            run_link = p_ev.add_run(f"[{desc}]\n")
                            run_link.font.underline = True
                            run_link.font.color.rgb = RGBColor(0, 89, 78) # Verde
                            run_link.font.bold = True
                            
                            who = f" (por: {ev.uploaded_by_name})" if ev.uploaded_by_name else ""
                            run_sub = p_ev.add_run(f"{url}{who}\n\n")
                            run_sub.font.size = Pt(7.5)
                            run_sub.font.color.rgb = RGBColor(100, 116, 139)
                    else:
                        p_ev.text = "Sin evidencias."
                        
                    # Zebra striping
                    row_bg = "F8FAFC" if r_idx % 2 == 0 else "FFFFFF"
                    for cell in row:
                        WordExporterService._set_cell_background(cell, row_bg)
                        WordExporterService._set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.font.size = Pt(8.5)
                                
                    row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                task_col_widths = [Inches(1.8), Inches(1.0), Inches(0.8), Inches(1.8), Inches(1.1)]
                for r in table.rows:
                    for c_idx, width in enumerate(task_col_widths):
                        r.cells[c_idx].width = width
                        
                doc.add_paragraph("\n")

        # --- SECCIÓN 4: GESTIÓN DE RECURSOS ---
        WordExporterService._add_styled_heading(doc, "4. GESTIÓN DE RECURSOS", level=1, space_before=12, space_after=12)
        doc.add_paragraph(
            "La ejecución de los recursos asignados a la Oficina de Talento Humano se desarrolló con estricto apego al "
            "Plan Operativo Anual de Inversiones (POAI), garantizando la viabilidad financiera de las actividades de "
            "bienestar social laboral, capacitación del personal y fortalecimiento de los sistemas mínimos de SST."
        )

        # --- SECCIÓN 5: ACCIONES DE MEJORA PARA IMPLEMENTAR ---
        WordExporterService._add_styled_heading(doc, "5. ACCIONES DE MEJORA PARA IMPLEMENTAR", level=1, space_before=12, space_after=12)
        doc.add_paragraph("A partir de las evaluaciones y el avance de las tareas, se proponen las siguientes acciones de mejora:")
        mejora_items = [
            "Avanzar en el proceso de adquisición e implementación del software digital integrado de Talento Humano (Novasoft) para automatizar la evaluación de desempeño.",
            "Promover acciones afirmativas tempranas para el cumplimiento de las cuotas de personas con discapacidad en futuras vacancias.",
            "Consolidar la Tiendita de la Honestidad y otras estrategias de la política de integridad como actividades permanentes en la institución.",
            "Realizar un monitoreo regular sobre la justificación de retrasos de tareas críticas para mitigar riesgos de incumplimiento."
        ]
        for item in mejora_items:
            doc.add_paragraph(item, style='List Bullet')

        # --- SECCIÓN 6: FIRMAS DE APROBACIÓN ---
        doc.add_paragraph("\n" * 2)
        WordExporterService._add_styled_heading(doc, "6. FIRMAS DE APROBACIÓN", level=1, space_before=12, space_after=12)
        
        table_signatures = doc.add_table(rows=1, cols=2)
        table_signatures.style = 'Table Grid'
        
        cells = table_signatures.rows[0].cells
        
        # Columna 1
        p1 = cells[0].paragraphs[0]
        p1.add_run("Elaboró / Reportó:\n\n\n\n___________________________________\n").font.bold = True
        p1.add_run("LILIA ANDREA NOCUA NEME\n").font.bold = True
        p1.add_run("Jefe Oficina de Talento Humano\nUnitrópico")
        
        # Columna 2
        p2 = cells[1].paragraphs[0]
        p2.add_run("Revisó / Aprobó:\n\n\n\n___________________________________\n").font.bold = True
        p2.add_run("DIRECCIÓN ADMINISTRATIVA / RECTORÍA\n").font.bold = True
        p2.add_run("Universidad Internacional del Trópico Americano")
        
        for c in cells:
            WordExporterService._set_cell_background(c, "F8FAFC")
            WordExporterService._set_cell_margins(c, top=150, bottom=150, left=150, right=150)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    r.font.size = Pt(9.5)
                    
        for r in table_signatures.rows:
            r.cells[0].width = Inches(3.25)
            r.cells[1].width = Inches(3.25)

        # Guardar en memoria
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio
